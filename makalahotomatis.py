import os
import win32com.client  # Untuk konversi ke PDF (Windows)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MakalahWord:
    def __init__(self):
        self.doc = Document()

    def set_font(self, paragraph, size, bold=False):
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)

    def add_cover(self, judul, nama_penulis, nim, prodi, institusi, tahun, logo_path):
        self.doc.add_paragraph()
        cover_section = self.doc.add_paragraph(judul.upper(), style='Heading 1')
        cover_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(cover_section, 18, True)
        for _ in range(5):
            self.doc.add_paragraph()

        if os.path.exists(logo_path):
            try:
                run = self.doc.add_paragraph().add_run()
                run.add_picture(logo_path, width=Inches(3.0))
                logo_paragraph = self.doc.paragraphs[-1]
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error saat menambahkan logo: {e}")
        else:
            print(f"File logo tidak ditemukan di {logo_path}")

        for _ in range(5):
            self.doc.add_paragraph()

        self.doc.add_paragraph(f"Nama Penulis: {nama_penulis}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"NIM: {nim}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"Program Studi: {prodi.upper()}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"Institusi: {institusi.upper()}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"Tahun: {tahun}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

        for para in self.doc.paragraphs[-5:]:
            self.set_font(para, 12)

        self.doc.add_paragraph()

    def add_kata_pengantar(self, kata_pengantar):
        self.doc.add_paragraph()
        kata_pengantar_judul = self.doc.add_paragraph('KATA PENGANTAR', style='Heading 1')
        kata_pengantar_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(kata_pengantar_judul, 14, True)
        self.doc.add_paragraph()
        kata_pengantar_paragraph = self.doc.add_paragraph(kata_pengantar, style='Normal')
        kata_pengantar_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.set_font(kata_pengantar_paragraph, 12, False)
        self.doc.add_page_break()

    def add_bab(self, judul_bab, subjudul=None, isi_subbab=None):
        bab_title = self.doc.add_paragraph(judul_bab, style='Heading 1')
        bab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(bab_title, 14, True)

        self.doc.add_paragraph()  # Menambahkan spasi setelah judul bab

        for sub, isi in zip(subjudul, isi_subbab):
            sub_title = self.doc.add_paragraph(sub, style='Heading 2')
            sub_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.set_font(sub_title, 12, True)

            paragraph = self.doc.add_paragraph(isi.strip(), style='Normal')

            # Tentukan alignment berdasarkan subjudul
            if "Rumusan Masalah" in sub:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Perataan kiri untuk Rumusan Masalah
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify untuk teks lain

            self.set_font(paragraph, 12)
            self.doc.add_page_break()

    def add_landasan_teori(self, jumlah_teori, teori_list):
        self.doc.add_paragraph("BAB 2: LANDASAN TEORI", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(self.doc.paragraphs[-1], 14, True)

        self.doc.add_paragraph()  # Menambahkan spasi setelah judul bab

        for i in range(jumlah_teori):
            self.doc.add_paragraph(f"2.{i+1} Teori {i+1}", style='Heading 2')
            self.set_font(self.doc.paragraphs[-1], 12, True)
            self.doc.add_paragraph(teori_list[i], style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.set_font(self.doc.paragraphs[-1], 12)
            self.doc.add_page_break()

            self.doc.add_paragraph()  # Menambahkan spasi setelah judul bab

    def add_penutup(self, saran, kesimpulan):
        self.doc.add_paragraph("BAB 3: PENUTUP", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(self.doc.paragraphs[-1], 14, True)

        self.doc.add_paragraph()  # Menambahkan spasi setelah judul bab

        self.doc.add_paragraph("3.1 Saran", style='Heading 2')
        self.set_font(self.doc.paragraphs[-1], 12, True)
        self.doc.add_paragraph(saran, style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.set_font(self.doc.paragraphs[-1], 12)
        self.doc.add_page_break()

        self.doc.add_paragraph("3.2 Kesimpulan", style='Heading 2')
        self.set_font(self.doc.paragraphs[-1], 12, True)
        self.doc.add_paragraph(kesimpulan, style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.set_font(self.doc.paragraphs[-1], 12)
        self.doc.add_page_break()

    def add_referensi(self, daftar_pustaka):
        judul_daftar_pustaka = self.doc.add_paragraph('DAFTAR PUSTAKA', style='Heading 1')
        judul_daftar_pustaka.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(judul_daftar_pustaka, 14, True)
        self.doc.add_paragraph()
        for item in daftar_pustaka:
            paragraph = self.doc.add_paragraph(item, style='Normal')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.set_font(paragraph, 12, False)

    def add_footer(self):
        try:
            temp_filename = os.path.join(os.getcwd(), "temp_makalah.docx")
            self.save(temp_filename)

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(temp_filename))

            # Menambahkan footer dan nomor halaman pada setiap section
            for section in doc.Sections:
                footer = section.Footers(2)  # Menentukan footer bawah
                page_numbers = footer.PageNumbers
                page_numbers.RestartNumberingAtSection = True  # Restart nomor halaman di setiap section
                page_numbers.Add(3, True)  # Menambahkan nomor halaman di footer (3: posisi di footer, True: nomor halaman di footer)
                page_numbers.Format = 4  # Format "Page X of Y"
                footer.Range.ParagraphFormat.Alignment = 1  # Meratakan nomor halaman ke tengah

            # Menghapus nomor halaman pada halaman pertama (cover)
            doc.Sections(1).Footers(2).PageNumbers.RestartNumberingAtSection = False
            doc.Sections(1).Footers(2).PageNumbers.Add(3, True)  # Menambahkan nomor halaman di footer
            doc.Sections(1).Footers(2).PageNumbers.RestartNumberingAtSection = False  # Hapus nomor halaman dari cover

            doc.SaveAs(os.path.abspath(temp_filename))
            doc.Close()
            word.Quit()

            pdf_filename = temp_filename.replace('.docx', '.pdf')
            self.convert_to_pdf(temp_filename)

            print(f"Makalah berhasil dikonversi ke PDF dengan footer: {pdf_filename}")
            os.remove(temp_filename)

        except Exception as e:
            print(f"{e}")

    def save(self, filename):
        self.doc.save(filename)

    def convert_to_pdf(self, filename):
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(os.path.abspath(filename))
        pdf_filename = filename.replace('.docx', '.pdf')
        doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)
        doc.Close()
        word.Quit()
        return pdf_filename


def buat_makalah():
    print("\n=== Buat Makalah Otomatis ===")

    judul = input("Masukkan judul makalah: ").strip()
    nama_penulis = input("Masukkan nama penyusun: ").strip()
    nim = input("Masukkan NIM: ").strip()
    prodi = input("Masukkan program studi: ").strip()
    institusi = input("Masukkan institusi: ").strip()
    tahun = input("Masukkan tahun: ").strip()

    logo_path = r"C:\Users\ARGA\Pictures\\254721151_utb_kotak.png"
    kata_pengantar = input("Masukkan Kata Pengantar: ").strip()

    subjudul_bab1 = ["1.1 Latar Belakang", "1.2 Rumusan Masalah", "1.3 Tujuan"]
    latar_belakang = input("Masukkan Latar Belakang: ").strip()

    rumusan_masalah = []
    jumlah_masalah = int(input("Berapa jumlah rumusan masalah? "))
    for i in range(jumlah_masalah):
        masalah = input(f"Masukkan Rumusan Masalah {i+1}: ").strip()
        rumusan_masalah.append(f"{i+1}. {masalah}")

    tujuan_makalah = input("Masukkan Tujuan: ").strip()

    jumlah_teori = int(input("Berapa jumlah teori yang ingin dimasukkan (maksimal 10)? "))
    teori_list = []
    for i in range(jumlah_teori):
        teori = input(f"Masukkan Teori {i+1}: ").strip()
        teori_list.append(teori)

    saran = input("Masukkan Saran: ").strip()
    kesimpulan = input("Masukkan Kesimpulan: ").strip()

    daftar_pustaka = []
    jumlah_referensi = int(input("Berapa jumlah referensi? "))
    for _ in range(jumlah_referensi):
        referensi = input("Masukkan referensi: ").strip()
        daftar_pustaka.append(referensi)

    makalah = MakalahWord()
    makalah.add_cover(judul, nama_penulis, nim, prodi, institusi, tahun, logo_path)
    makalah.add_kata_pengantar(kata_pengantar)
    makalah.add_bab("BAB 1: Pendahuluan", subjudul_bab1, [latar_belakang, '\n'.join(rumusan_masalah), tujuan_makalah])
    makalah.add_landasan_teori(jumlah_teori, teori_list)
    makalah.add_penutup(saran, kesimpulan)
    makalah.add_referensi(daftar_pustaka)
    makalah.add_footer()

    file_name = f"{judul.replace(' ', '_')}_makalah.docx"
    makalah.save(file_name)
    print(f"Makalah berhasil dibuat: {file_name}")

if __name__ == "__main__":
    buat_makalah()